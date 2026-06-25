import io
from flask import jsonify, send_file, request
from docx import Document
import openpyxl
from app.models.case import Case
from app.models.transaction import Transaction
from app.reports.report_generator import reports_bp

@reports_bp.route('/dossier/<case_id>', methods=['GET'])
def generate_dossier(case_id):
    authority = request.args.get('authority', 'bank_fraud')
    case = Case.query.get(case_id)
    
    if not case:
        return jsonify({"error": "Case not found"}), 404
        
    transactions = Transaction.query.filter_by(case_id=case_id, is_failed=False).all()
    
    if authority == 'auditor':
        # Generate Excel
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Raw Data"
        headers = ["Date", "Amount", "Type", "Sender", "Receiver", "Description", "Balance After"]
        ws.append(headers)
        
        for t in transactions:
            ws.append([
                t.date.strftime('%Y-%m-%d') if t.date else "",
                t.amount,
                t.type,
                t.sender_account,
                t.receiver_account,
                t.description,
                t.balance_after
            ])
            
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"Dossier_Auditor_{case_id}.xlsx",
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        
    else:
        # Generate DOCX
        doc = Document()
        doc.add_heading(f"FINTELLIGENCE Dossier: {authority.replace('_', ' ').title()}", 0)
        
        doc.add_heading("Case Details", level=1)
        doc.add_paragraph(f"Case ID: {case_id}")
        doc.add_paragraph(f"Risk Score: {case.suspicion_score}")
        
        if authority == 'bank_fraud':
            doc.add_heading("Transaction Details & Pass-Through", level=1)
            doc.add_paragraph("This section contains transaction amounts, account details, and pass-through ratios relevant for bank fraud investigation.")
            # Simplified mock content
            
        elif authority == 'aml_team':
            doc.add_heading("AML Patterns & Layering Chains", level=1)
            doc.add_paragraph("This section details identified AML patterns such as layering chains and structuring evidence.")
            
        elif authority == 'cyber_crime':
            doc.add_heading("Digital Trail & Timing Analysis", level=1)
            doc.add_paragraph("This section focuses on the digital transaction trail and odd-hour timing anomalies.")
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"Dossier_{authority}_{case_id}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
